#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Generator Service
Generates Word documents with translations
"""

import os
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocumentGenerator:
    """Generate Word documents with translations."""

    def __init__(self):
        pass

    def set_chinese_font(self, run, font_name: str = "SimSun", font_size: int = 11, bold: bool = False):
        """Set Chinese font for a run."""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def set_cell_border(self, cell, **kwargs):
        """Set cell border."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'
                element = OxmlElement(tag)
                element.set(qn('w:val'), edge_data.get('val', 'single'))
                element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                element.set(qn('w:color'), edge_data.get('color', '000000'))
                tcBorders.append(element)

        tcPr.append(tcBorders)

    def generate_document(self, translations: List[Dict], output_path: str,
                         source_lang: str = 'zh', target_lang: str = 'en',
                         merge: bool = True) -> str:
        """
        Generate Word document with translations.

        Args:
            translations: List of translation results
            output_path: Path to save the document
            source_lang: Source language code
            target_lang: Target language code
            merge: Whether to merge into single document

        Returns:
            Path to generated document
        """
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        # Add document title
        self._add_title_page(doc, source_lang, target_lang)

        # Add table of contents placeholder
        self._add_toc(doc, translations)

        # Add page break
        doc.add_page_break()

        # Add each translation section
        for i, trans in enumerate(translations):
            self._add_section(doc, trans, i + 1, source_lang, target_lang)

            # Add page break between sections (except last)
            if i < len(translations) - 1:
                doc.add_page_break()

        # Add certification page
        self._add_certification(doc)

        # Save document
        doc.save(output_path)
        return output_path

    def _add_title_page(self, doc: Document, source_lang: str, target_lang: str):
        """Add title page."""
        lang_names = {
            'zh': '中文', 'en': 'English', 'ja': '日本語', 'ko': '한국어',
            'fr': 'Français', 'de': 'Deutsch', 'es': 'Español', 'ru': 'Русский'
        }

        # Main title
        title = doc.add_heading('签证材料翻译件', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('Visa Document Translation', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Language info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info_para.add_run(f"原文语言 / Source Language: {lang_names.get(source_lang, source_lang)}")
        self.set_chinese_font(run, "SimSun", 12)

        info_para2 = doc.add_paragraph()
        info_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = info_para2.add_run(f"译文语言 / Target Language: {lang_names.get(target_lang, target_lang)}")
        self.set_chinese_font(run2, "SimSun", 12)

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = datetime.now().strftime('%Y年%m月%d日 / %B %d, %Y')
        run = date_para.add_run(f"生成日期 / Generated: {date_str}")
        self.set_chinese_font(run, "SimSun", 10)

    def _add_toc(self, doc: Document, translations: List[Dict]):
        """Add table of contents."""
        toc_title = doc.add_heading('目录 / Table of Contents', level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Create TOC table
        toc_table = doc.add_table(rows=len(translations) + 1, cols=3)
        toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_row = toc_table.rows[0]
        headers = ['序号 / No.', '文档名称 / Document', '类型 / Type']
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header_text)
            self.set_chinese_font(run, "SimHei", 10, bold=True)

        # TOC entries
        for i, trans in enumerate(translations):
            row = toc_table.rows[i + 1]

            # Number
            row.cells[0].paragraphs[0].add_run(str(i + 1))
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Document name
            image_path = trans.get('image_path', '')
            filename = Path(image_path).name if image_path else f'Document {i+1}'
            run = row.cells[1].paragraphs[0].add_run(filename)
            self.set_chinese_font(run, "SimSun", 10)

            # Document type
            doc_type = trans.get('translation', {}).get('document_type', 'Unknown')
            run = row.cells[2].paragraphs[0].add_run(doc_type)
            self.set_chinese_font(run, "SimSun", 10)

    def _add_section(self, doc: Document, trans: Dict, section_num: int,
                    source_lang: str, target_lang: str):
        """Add a translation section."""
        image_path = trans.get('image_path', '')
        translation = trans.get('translation', {})

        # Section title
        filename = Path(image_path).name if image_path else f'Document {section_num}'
        section_title = doc.add_heading(f'{section_num}. {filename}', level=1)
        section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Document type subtitle
        doc_type = translation.get('document_type', 'Unknown Document')
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = type_para.add_run(f'文档类型 / Document Type: {doc_type}')
        self.set_chinese_font(run, "SimSun", 11, bold=True)

        doc.add_paragraph()

        # Add original image
        if image_path and Path(image_path).exists():
            orig_title = doc.add_heading('原文图片 / Original Document', level=2)

            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=Inches(5.5))

                # Caption
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(f'图 {section_num}: {filename}')
                self.set_chinese_font(run, "SimSun", 10)

                doc.add_paragraph()
            except Exception as e:
                error_para = doc.add_paragraph()
                error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = error_para.add_run(f'[图片加载失败 / Image load failed: {str(e)}]')
                self.set_chinese_font(run, "SimSun", 10)

        # Add translation table
        trans_title = doc.add_heading('翻译对照表 / Translation Table', level=2)

        fields = translation.get('fields', [])
        if fields:
            self._add_translation_table(doc, fields)
        else:
            # If no structured fields, show full text
            full_original = translation.get('full_text_original', '')
            full_translated = translation.get('full_text_translated', '')

            if full_translated:
                para = doc.add_paragraph()
                run = para.add_run('译文 / Translation:')
                self.set_chinese_font(run, "SimHei", 11, bold=True)

                trans_para = doc.add_paragraph()
                run = trans_para.add_run(full_translated)
                self.set_chinese_font(run, "SimSun", 10)
            else:
                raw = translation.get('raw_response', '翻译失败 / Translation failed')
                para = doc.add_paragraph()
                run = para.add_run(raw)
                self.set_chinese_font(run, "SimSun", 10)

    def _add_translation_table(self, doc: Document, fields: List[Dict]):
        """Add translation comparison table."""
        if not fields:
            return

        # Create table
        table = doc.add_table(rows=len(fields) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.5)  # Original label
            row.cells[1].width = Inches(2.0)  # Original value
            row.cells[2].width = Inches(1.5)  # Translated label
            row.cells[3].width = Inches(2.0)  # Translated value

        # Set borders
        border_style = {'val': 'single', 'sz': 6, 'color': '000000'}
        for row in table.rows:
            for cell in row.cells:
                self.set_cell_border(cell, top=border_style, left=border_style,
                                    bottom=border_style, right=border_style)

        # Header row
        headers = ['原文项目', '原文内容', '译文项目', '译文内容']
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header_text)
            self.set_chinese_font(run, "SimHei", 9, bold=True)

        # Data rows
        for i, field in enumerate(fields):
            row = table.rows[i + 1]

            # Original label
            label_orig = field.get('label_original', '')
            run = row.cells[0].paragraphs[0].add_run(label_orig)
            self.set_chinese_font(run, "SimSun", 9, bold=True)

            # Original value
            value_orig = field.get('value_original', '')
            run = row.cells[1].paragraphs[0].add_run(value_orig)
            self.set_chinese_font(run, "SimSun", 9)

            # Translated label
            label_trans = field.get('label_translated', '')
            run = row.cells[2].paragraphs[0].add_run(label_trans)
            run.font.size = Pt(9)
            run.font.name = "Arial"

            # Translated value
            value_trans = field.get('value_translated', '')
            run = row.cells[3].paragraphs[0].add_run(value_trans)
            run.font.size = Pt(9)
            run.font.name = "Arial"

    def _add_certification(self, doc: Document):
        """Add certification page."""
        doc.add_page_break()

        cert_title = doc.add_heading('翻译声明 / Translation Certification', level=1)
        cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Certification text
        cert_text = """本人声明：本翻译件中的所有翻译内容均与原件一致，真实有效。
翻译件仅供签证申请参考使用，原件应与翻译件一并提交。

I hereby certify that all translations in this document are true and accurate
to the best of my knowledge. This translation is for visa application reference only.
The original documents should be submitted together with this translation."""

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(cert_text)
        self.set_chinese_font(run, "SimSun", 11)

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature area
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = sign_para.add_run('翻译人签字 / Translator Signature: ________________\n\n')
        run.font.size = Pt(11)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_para.add_run('日期 / Date: ________________')
        run.font.size = Pt(11)
